#!/usr/bin/env python
# coding: utf-8

# In[1]:


#imports
import pandas as pd
from datetime import date
import docx as dx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import os
from sqlalchemy import create_engine, MetaData, Table, Column, String, DateTime, insert, delete


# In[12]:


#connect to database containing previously applied jobs (or create db)
engine = create_engine('sqlite:///applications.db') #make db file in cwd
metadata = MetaData()

applications = Table('applications', metadata,
                Column('position', String(255)),
                Column('company', String(255)),
                Column('date', DateTime()))

metadata.create_all(engine)


# In[ ]:


#see 5 most recent job applications
q = 'SELECT * FROM applications ORDER BY date DESC LIMIT 5'
pd.read_sql(q, engine)


# In[3]:


# insert coverletters
default_cover_letter = '''

YOUR LETTER HERE
Use {d} for where the date goes, {p} for position youre applying for, and {c} for company name
See example below

'''

example_cl = '''Harry J Potter
The Cupboard under the Stairs
4, Privet Drive
Little Whinging
Surry


{d}
Application for {p}

Dear Hiring Manager,

I’m an auror and I’d love to bring my courage and experience to your team. I am looking for work that continues to expand my magical ability, and I believe the {p} position at {c} will provide challenges that will help me grow as a dark wizard catcher and as a team player.

I have previously defeated Lord Voldemort. Through this experience I learned the stamina and interpersonal communication required to find and destroy each of his seven horcruxes.

I’m eager to speak more in depth with you about the position. Please send an owl to get in touch.

Thank you for your time and consideration,

Harry Potter'''


cover_letters = [default_cover_letter, example_cl]


# In[4]:


#find info to fill in the blanks d, p, and c
today = date.today().strftime("%m/%d/%Y") #find date

position = input('Position Name: ') #position name
company = input('Company Name: ') #company name
cl_select = input('Which coverletter? 0: default 1: example  ') #which cover letter to use


# In[5]:


#choose coverletter and fill in blanks
complete = cover_letters[int(cl_select)].format(d = today, p = position, c = company)
print(complete)


# In[23]:


#split coverletter into pieces for word to format
clsplit = complete.split('\n') #split on breaks
while '' in clsplit: #remove lines with empty strings
    clsplit.remove('')
clsplit.insert(7, '') #put an empty string back in to make the letter formating nice.

#create word doc
doc = dx.Document()
name = doc.add_paragraph('') #bold and right align name
name.add_run(clsplit.pop(0)).bold = True
name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for _ in range(4): #for the rest of the contact info right align
    line = doc.add_paragraph(clsplit.pop(0))
    line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for _ in range(len(clsplit)): #rest of the doc left align
    line = doc.add_paragraph(clsplit.pop(0))
word_filename = company+'_coverletter.docx' 
doc.save(word_filename)

pdf_filename = word_filename.replace('.docx','.pdf')
convert(word_filename,pdf_filename) #convert word to pdf

os.remove(word_filename) #remove word file

#create new entry in database table
new = insert(applications).values(position = position, company = company,
                                  date = date.today())
engine.execute(new)


# In[ ]:


q = 'SELECT * FROM applications' #see all job apps saved
pd.read_sql(q, engine) #wow so productive!


# In[ ]:




